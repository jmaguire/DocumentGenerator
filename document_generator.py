import json
import re
import logging
from typing import Any, Dict
from functools import lru_cache

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import html2text


class DocumentGenerator:
    BORDER_COLOR_DEFAULT = "000000"
    BORDER_SIZE_DEFAULT = "4"
    CELL_BG_COLOR_WHITE = "FFFFFF"
    CELL_BG_COLOR_HEADER = "94B1D4"
    CELL_BG_COLOR_BLUE = "F6FAFF"
    TABLE_STYLE = 'Table Grid'
    BORDER_COLOR_GREY = "CCCCCC"
    CELL_BG_COLOR_DARK = "002541"

    def __init__(self, data_file: str):
        self.data_file = data_file
        self.document = Document()
        self.h = html2text.HTML2Text()
        self.h.body_width = 0
        self.setup_logging()

    @staticmethod
    def setup_logging():
        logging.basicConfig(level=logging.INFO,
                            format='%(asctime)s - %(levelname)s - %(message)s')

    @lru_cache(maxsize=None)
    def clean_text(self, text: str) -> str:
        """
        Clean and format text from HTML to plain text.

        :param text: The HTML text to clean.
        :return: The cleaned plain text.
        """
        pattern = r'\n[ \t]*\n[ \t]*\n'
        text = self.h.handle(text).strip()
        return re.sub(pattern, '\n\n', text)

    def parse_markdown_table(self, markdown):
        lines = markdown.strip().split('\n')
        headers = re.split(r'\s*\|\s*', lines[0].strip())
        rows = [re.split(r'\s*\|\s*', line.strip()) for line in lines[2:]]
        return headers, rows

    def handle_answer_type(self, answer_value: str, answer_type: str) -> str:
        if answer_type == "Currency":
            return answer_value.replace(" ", "")
        if answer_type == "Percentage":
            return answer_value.strip() + "%"
        if answer_type == "Multiple Choice(Select all that apply)":
            return answer_value.strip()
        return self.clean_text(answer_value)

    def set_table_border(self, table, border_color=BORDER_COLOR_DEFAULT, border_size=BORDER_SIZE_DEFAULT):
        tbl = table._element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), border_size)
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)

    def set_cell_background(self, cell, color):
        cell_properties = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        cell_properties.append(shd)

    def set_cell_margins(self, cell, top=0.08, start=0.16, bottom=0.08, end=0.16):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_type, margin_size in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            mar = OxmlElement(f'w:{margin_type}')
            mar.set(qn('w:w'), str(int(margin_size * 1440)))
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        tcPr.append(tcMar)

    def set_cell_width(self, cell, width):
        cell.width = Inches(width)

    def _style_table_row(self, row_cells, header_color=CELL_BG_COLOR_BLUE, body_color=CELL_BG_COLOR_WHITE, width=.75):
        self.set_cell_background(row_cells[0], header_color)
        self.set_cell_margins(row_cells[0], .04, .08, .04, .08)
        self.set_cell_width(row_cells[0], width)
        self.set_cell_background(row_cells[1], body_color)
        self.set_cell_margins(row_cells[1], .04, .08, .04, .08)

    def _add_question_header(self, question_table, quest: Dict[str, Any]):
        row_cells = question_table.add_row().cells
        row_cells[0].merge(row_cells[1])
        row_cells[0].text = f"{self.clean_text(quest['elementNumber'])} {
            self.clean_text(quest['questionText'])}"
        self.set_cell_background(row_cells[0], self.CELL_BG_COLOR_HEADER)
        self.set_cell_margins(row_cells[0], .08, .08, .08, .08)

    def _create_question_table(self):
        question_table = self.document.add_table(
            rows=0, cols=2, style=self.TABLE_STYLE)
        self.set_table_border(
            question_table, border_color=self.BORDER_COLOR_GREY, border_size=self.BORDER_SIZE_DEFAULT)
        return question_table

    def _populate_table(self, table, headers: list, rows: list):
        hdr_cells = table.add_row().cells
        for index, header in enumerate(headers):
            hdr_cells[index].text = header
            self.set_cell_background(
                hdr_cells[index], self.CELL_BG_COLOR_WHITE)
        for row in rows:
            data_cells = table.add_row().cells
            for index, value in enumerate(row):
                data_cells[index].text = value
                self.set_cell_background(
                    data_cells[index], self.CELL_BG_COLOR_WHITE)

    def build_question(self, quest):
        def add_question_row(table, name, value):
            value = value if value else "N/A"
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = value
            if value == "N/A":
                row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(124, 124, 124)
            self._style_table_row(row_cells)
            return table

        def add_question_table_row(table, name, value):
            headers, rows = self.parse_markdown_table(value)
            row_cells = table.add_row().cells
            row_cells[0].text = name
            table_answer = row_cells[1].add_table(rows=0, cols=len(headers))
            self.set_table_border(
                table_answer, border_color=self.BORDER_COLOR_GREY, border_size=self.BORDER_SIZE_DEFAULT)
            self._populate_table(table_answer, headers, rows)
            self._style_table_row(row_cells)
            return table

        question_table = self.document.add_table(
            rows=0, cols=2, style=self.TABLE_STYLE)
        self.set_table_border(
            question_table, border_color=self.BORDER_COLOR_GREY, border_size=self.BORDER_SIZE_DEFAULT)

        question_table = self._create_question_table()
        self._add_question_header(question_table, quest)

        # Handle followup
        if quest["triggerValue"]:
            followup_text = self.clean_text(
                quest["triggerValue"]) + " on " + self.clean_text(quest["parentElementNumber"])
            question_table = add_question_row(
                question_table, "Follow-up", followup_text)

        # Add answer
        answer_value = self.handle_answer_type(
            quest["answerValue"], quest["answerType"])
        if quest["answerType"] != "Table" and quest["answerType"] != "JSON":
            question_table = add_question_row(
                question_table, "Answer", answer_value)
        else:
            question_table = add_question_table_row(
                question_table, "Answer", answer_value)

        # Comments and documents
        question_table = add_question_row(
            question_table, "Comment", self.clean_text(quest["answerComments"]))
        documents = [doc for docList in quest["docListAnswer"]
                     for doc in docList["attachedFiles"]]
        documentAsList = ", ".join(documents)
        question_table = add_question_row(
            question_table, "Documents", documentAsList)

    def get_nested_value(self, data, keys, default=None):
        for key in keys:
            if isinstance(data, dict):
                data = data.get(key, default)
            else:
                return default
            if data is default:
                break
        return data

    def add_meta_data_row(self, table, name, value):
        if not value:
            return table
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = value
        self._style_table_row(row_cells, header_color=self.CELL_BG_COLOR_HEADER, width=1.5)
        return table

    def add_assessment_metadata(self, data: Dict[str, Any]):
        """Add the section details to the document."""
        table = self.document.add_table(rows=0, cols=2, style=self.TABLE_STYLE)
        self.set_table_border(
            table, border_color=self.BORDER_COLOR_GREY, border_size=self.BORDER_SIZE_DEFAULT)

        assessment_details = [
            ("Assessment", [
             "pqiBasicInfo", "questionnaireBasicInfo", "qStructureBasicInfo", "name"]),
            ("Report Generated", ["reportGeneratedOn"]),
            ("Publisher", ["pqiBasicInfo", "publishedByUser", "companyName"]),
            ("Published By", ["pqiBasicInfo",
             "publishedByUser", "userProfile", "fullName"]),
            ("Published On", ["pqiBasicInfo", "publishedDate"]),
            ("Recipient", ["partner", "name"]),
            ("Received By", ["pqiBasicInfo",
             "publishedToContact", "contactProfile", "fullName"])
        ]

        for detail_name, detail_keys in assessment_details:
            detail_value = self.get_nested_value(data, detail_keys, "")
            self.add_meta_data_row(table, detail_name, detail_value)

    def set_page_margins(self, margin_lr=Inches(1), margin_tb=Inches(0.5)):
        for section in self.document.sections:
            section.top_margin = margin_tb
            section.bottom_margin = margin_tb
            section.left_margin = margin_lr
            section.right_margin = margin_lr

    def add_assessment_section(self, sectionHeading):
        """Add the section or subsection name to the document."""
        table = self.document.add_table(rows=0, cols=1, style=self.TABLE_STYLE)
        self.set_table_border(
            table, border_color=self.BORDER_COLOR_GREY, border_size=self.BORDER_SIZE_DEFAULT)
        row_cells = table.add_row().cells
        row_cells[0].text = sectionHeading
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        self.set_cell_background(row_cells[0], self.CELL_BG_COLOR_DARK)
        self.set_cell_margins(row_cells[0], .04, .08, .04, .08)

    def generate_document(self):

        try:
            with open(self.data_file, "r") as f:
                data = json.load(f)
        except FileNotFoundError:
            logging.error("Data file not found.")
            return
        except json.JSONDecodeError:
            logging.error("Error decoding JSON from the data file.")
            return
        
        ## Define document properties
        # Update paragraph font
        self.set_page_margins(Inches(.5), Inches(.5))
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Aptos'
        font.size = Pt(11)

        # Add custom header
        styles = self.document.styles
        new_heading_style = styles.add_style('Centrl Header', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 1']
        font = new_heading_style.font
        font.name = 'Aptos'
        font.size = Pt(16)
        font.bold = False 
        self.document.add_paragraph('Assessment Export', style='Centrl Header')

        # Add assessment metadata
        self.add_assessment_section("Details")
        self.add_assessment_metadata(data)
        self.document.add_paragraph()

        ## Add question information
        for section in data["sections"]:
            section_heading = f"{self.clean_text(section['elementNumber'])}. {self.clean_text(section['sectionName'])}"
            self.add_assessment_section(section_heading)
            #self.document.add_heading(section_heading, level=2)
            for quest in section["questionDetails"]:
                self.build_question(quest)
            
            for subSection in section["subSections"]:
                sub_section_heading = f"{self.clean_text(subSection['elementNumber'])}. {self.clean_text(subSection['sectionName'])}"
                self.add_assessment_section(sub_section_heading)
                for quest in subSection["questionDetails"]:
                    self.build_question(quest)
                self.document.add_paragraph()
            self.document.add_paragraph()


        self.document.save('demo.docx')
        logging.info("Document generated successfully.")


if __name__ == "__main__":
    doc_gen = DocumentGenerator("afme.json")
    try:
        doc_gen.generate_document()
    except Exception as e:
        logging.error(f"Error generating document: {e}")
